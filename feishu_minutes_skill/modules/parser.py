"""
会议纪要解析模块
支持解析飞书妙记链接和 docx 文件
"""

import re
from dataclasses import dataclass
from typing import List, Optional
from datetime import datetime


@dataclass
class ActionItem:
    """行动项"""
    content: str
    assignee: Optional[str]  # 负责人
    deadline: Optional[str]  # 截止时间
    
    def has_assignee(self) -> bool:
        return self.assignee is not None and self.assignee.strip() != ""


@dataclass
class MeetingMinutes:
    """会议纪要数据结构"""
    title: str
    date: Optional[str]
    attendees: List[str]
    conclusions: List[str]
    action_items: List[ActionItem]
    raw_content: str


class MinutesParser:
    """纪要解析器"""
    
    # 匹配@人的模式
    MENTION_PATTERN = re.compile(r'@([^\s@,，。；;:\n]+)')
    # 匹配截止日期的模式
    DEADLINE_PATTERN = re.compile(r'(截止|DDL|deadline|期限)[：:\s]*(\d{1,2}[月/]\d{1,2}[日号]?|\d{4}-\d{2}-\d{2}|下周[一二三四五六日]|明天|后天)')
    # 匹配行动项的模式（以TODO、待办、行动项等开头）
    ACTION_ITEM_PATTERN = re.compile(r'^(?:[\d一二三四五六七八九十]+[.．、]|[-•*]|TODO|待办|行动项|Action)[\s]*', re.IGNORECASE)
    
    def parse_from_text(self, text: str, title: str = "会议纪要") -> MeetingMinutes:
        """
        从文本内容解析会议纪要
        
        Args:
            text: 纪要文本内容
            title: 会议标题
            
        Returns:
            MeetingMinutes 对象
        """
        lines = text.split('\n')
        
        # 提取参会人
        attendees = self._extract_attendees(text)
        
        # 提取日期
        date = self._extract_date(text)
        
        # 提取结论和行动项
        conclusions = []
        action_items = []
        
        in_conclusion_section = False
        in_action_section = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # 检测章节
            if self._is_conclusion_header(line):
                in_conclusion_section = True
                in_action_section = False
                continue
            elif self._is_action_header(line):
                in_conclusion_section = False
                in_action_section = True
                continue
            elif self._is_other_header(line):
                in_conclusion_section = False
                in_action_section = False
                continue
            
            # 提取内容
            if in_conclusion_section and line:
                # 清理列表标记
                clean_line = self._clean_list_marker(line)
                if clean_line:
                    conclusions.append(clean_line)
            elif in_action_section and line:
                action_item = self._parse_action_item(line)
                if action_item:
                    action_items.append(action_item)
            elif self.ACTION_ITEM_PATTERN.match(line):
                # 未明确分区时的行动项检测
                action_item = self._parse_action_item(line)
                if action_item:
                    action_items.append(action_item)
        
        # 如果没有明确分区，尝试智能提取
        if not conclusions and not action_items:
            conclusions, action_items = self._smart_extract(text)
        
        return MeetingMinutes(
            title=title,
            date=date,
            attendees=attendees,
            conclusions=conclusions,
            action_items=action_items,
            raw_content=text
        )
    
    def _extract_attendees(self, text: str) -> List[str]:
        """提取参会人列表"""
        attendees = []
        
        # 查找参会人/与会人/参与人章节
        patterns = [
            r'(?:参会人|与会人|参与人|出席人|参会人员)[：:\s]*\n?([^\n]+)',
            r'(?:参会人|与会人|参与人|出席人|参会人员)[：:\s]*\n?((?:[^\n]+\n?){1,5})',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                attendee_text = match.group(1)
                # 分割多个参会人
                parts = re.split(r'[,，、;；\s]+', attendee_text)
                for part in parts:
                    part = part.strip()
                    if part and len(part) < 20:  # 过滤过长的内容
                        attendees.append(part)
                break
        
        # 如果没有找到，尝试从@提取
        if not attendees:
            mentions = self.MENTION_PATTERN.findall(text)
            attendees = list(set(mentions))
        
        return attendees
    
    def _extract_date(self, text: str) -> Optional[str]:
        """提取会议日期"""
        patterns = [
            r'(\d{4}[年/-]\d{1,2}[月/-]\d{1,2}[日号]?)',
            r'(\d{1,2}[月/-]\d{1,2}[日号]?)',
            r'(?:会议时间|时间)[：:\s]*(\d{4}[年/-]\d{1,2}[月/-]\d{1,2})',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(1)
        
        return None
    
    def _is_conclusion_header(self, line: str) -> bool:
        """判断是否为结论章节标题"""
        keywords = ['结论', '决议', '决策', '共识', '达成一致', '关键结论', '会议结论']
        return any(kw in line for kw in keywords) and len(line) < 20
    
    def _is_action_header(self, line: str) -> bool:
        """判断是否为行动项章节标题"""
        keywords = ['行动项', '待办', 'TODO', '行动计划', '后续行动', '跟进事项', 'Action Item']
        return any(kw in line for kw in keywords) and len(line) < 20
    
    def _is_other_header(self, line: str) -> bool:
        """判断是否为其他章节标题（用于退出当前章节）"""
        keywords = ['背景', '议题', '讨论', '记录', '备注', '附件']
        return any(kw in line for kw in keywords) and len(line) < 20
    
    def _clean_list_marker(self, line: str) -> str:
        """清理列表标记"""
        # 移除常见的列表标记
        cleaned = re.sub(r'^[\d一二三四五六七八九十]+[.．、]\s*', '', line)
        cleaned = re.sub(r'^[-•*]\s*', '', cleaned)
        return cleaned.strip()
    
    def _parse_action_item(self, line: str) -> Optional[ActionItem]:
        """解析单行行动项"""
        # 清理列表标记
        content = self._clean_list_marker(line)
        if not content:
            return None
        
        # 提取负责人
        assignee = None
        mentions = self.MENTION_PATTERN.findall(content)
        if mentions:
            assignee = mentions[0]
        
        # 提取截止日期
        deadline = None
        deadline_match = self.DEADLINE_PATTERN.search(content)
        if deadline_match:
            deadline = deadline_match.group(2)
        
        return ActionItem(
            content=content,
            assignee=assignee,
            deadline=deadline
        )
    
    def _smart_extract(self, text: str) -> tuple:
        """智能提取结论和行动项（无明确分区时）"""
        conclusions = []
        action_items = []
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line or len(line) < 5:
                continue
            
            # 检测行动项特征
            action_keywords = ['完成', '跟进', '处理', '落实', '准备', '提交', '确认', '联系', '协调']
            has_action_keyword = any(kw in line for kw in action_keywords)
            has_mention = '@' in line
            has_deadline = self.DEADLINE_PATTERN.search(line) is not None
            
            if has_mention or (has_action_keyword and has_deadline):
                action_item = self._parse_action_item(line)
                if action_item:
                    action_items.append(action_item)
            elif len(line) > 10 and not line.startswith('http'):
                # 可能是结论
                conclusions.append(line)
        
        return conclusions, action_items


class DocxParser(MinutesParser):
    """Docx 文件解析器"""
    
    def parse(self, file_path: str) -> MeetingMinutes:
        """
        解析 docx 文件
        
        Args:
            file_path: docx 文件路径
            
        Returns:
            MeetingMinutes 对象
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")
        
        doc = Document(file_path)
        
        # 提取标题（第一个段落）
        title = "会议纪要"
        if doc.paragraphs:
            first_text = doc.paragraphs[0].text.strip()
            if first_text and len(first_text) < 50:
                title = first_text
        
        # 提取所有文本
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        return self.parse_from_text(full_text, title)


class FeishuMinutesParser(MinutesParser):
    """飞书妙记解析器"""
    
    def parse_from_url(self, url: str, api_token: Optional[str] = None) -> MeetingMinutes:
        """
        从飞书妙记链接解析
        
        Args:
            url: 妙记链接
            api_token: 飞书 API Token（可选）
            
        Returns:
            MeetingMinutes 对象
        """
        # 提取会议 ID
        meeting_id = self._extract_meeting_id(url)
        if not meeting_id:
            raise ValueError(f"无法从 URL 提取会议 ID: {url}")
        
        # 如果有 API Token，尝试调用 API 获取
        if api_token:
            return self._fetch_from_api(meeting_id, api_token)
        else:
            # 提示用户需要导出 docx 或提供 API Token
            raise ValueError(
                "请提供飞书 API Token 以获取妙记内容，"
                "或将妙记导出为 docx 文件后上传"
            )
    
    def _extract_meeting_id(self, url: str) -> Optional[str]:
        """从妙记链接提取会议 ID"""
        patterns = [
            r'meetings\.feishu\.cn/minutes/([a-zA-Z0-9_-]+)',
            r'feishu\.cn/minutes/([a-zA-Z0-9_-]+)',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        
        return None
    
    def _fetch_from_api(self, meeting_id: str, api_token: str) -> MeetingMinutes:
        """通过飞书 API 获取妙记内容"""
        import requests
        
        # 飞书 API 端点（需要根据实际 API 文档调整）
        url = f"https://open.feishu.cn/open-apis/minutes/v1/minutes/{meeting_id}"
        
        headers = {
            "Authorization": f"Bearer {api_token}",
            "Content-Type": "application/json"
        }
        
        response = requests.get(url, headers=headers)
        response.raise_for_status()
        
        data = response.json()
        
        # 解析 API 返回的数据（根据实际 API 响应结构调整）
        title = data.get("data", {}).get("title", "会议纪要")
        content = data.get("data", {}).get("content", "")
        
        return self.parse_from_text(content, title)


def parse_minutes(source: str, is_url: bool = False, is_docx: bool = False, api_token: Optional[str] = None) -> MeetingMinutes:
    """
    统一解析入口
    
    Args:
        source: 妙记链接、docx 文件路径或文本内容
        is_url: 是否为 URL
        is_docx: 是否为 docx 文件
        api_token: 飞书 API Token
        
    Returns:
        MeetingMinutes 对象
    """
    if is_url:
        parser = FeishuMinutesParser()
        return parser.parse_from_url(source, api_token)
    elif is_docx:
        parser = DocxParser()
        return parser.parse(source)
    else:
        parser = MinutesParser()
        return parser.parse_from_text(source)
